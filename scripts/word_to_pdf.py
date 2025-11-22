#!/usr/bin/env python3
"""
Word to PDF conversion script with multiple methods for perfect formatting preservation.
Tries methods in order of quality:
1. docx2pdf (requires Word/LibreOffice) - Perfect match
2. LibreOffice headless - Very good match, cross-platform
3. python-docx + reportlab - Fallback with good formatting
"""

import sys
import os
import subprocess
from pathlib import Path

# Method 1: Try docx2pdf (requires Word on Windows or LibreOffice)
def try_docx2pdf(word_path: str, pdf_path: str) -> bool:
    """Try conversion using docx2pdf library"""
    try:
        from docx2pdf import convert
        # docx2pdf convert function takes input and output paths
        # On Windows, it uses Word COM automation
        # On Linux/Mac, it uses LibreOffice
        convert(word_path, pdf_path)
        
        # Check if conversion was successful
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            print(f"SUCCESS: Converted using docx2pdf: {word_path} to {pdf_path}", file=sys.stderr)
            return True
        else:
            print(f"Warning: docx2pdf conversion completed but output file not found or empty", file=sys.stderr)
    except ImportError:
        pass  # Library not installed
    except Exception as e:
        print(f"Warning: docx2pdf failed: {e}", file=sys.stderr)
    return False

# Method 2: Try LibreOffice headless (cross-platform, excellent results)
def try_libreoffice(word_path: str, pdf_path: str) -> bool:
    """Try conversion using LibreOffice headless mode"""
    try:
        # Find LibreOffice executable
        libreoffice_cmds = []
        
        if sys.platform == 'win32':
            # Windows paths
            libreoffice_cmds = [
                r'C:\Program Files\LibreOffice\program\soffice.exe',
                r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
                'soffice.exe',
            ]
        else:
            # Unix-like systems
            libreoffice_cmds = [
                'libreoffice',
                '/usr/bin/libreoffice',
                '/usr/local/bin/libreoffice',
            ]
        
        libreoffice_cmd = None
        for cmd in libreoffice_cmds:
            try:
                result = subprocess.run(
                    [cmd, '--version'],
                    capture_output=True,
                    timeout=5,
                    text=True
                )
                if result.returncode == 0:
                    libreoffice_cmd = cmd
                    break
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue
        
        if not libreoffice_cmd:
            return False
        
        # Convert using LibreOffice headless
        output_dir = os.path.dirname(pdf_path) or '.'
        result = subprocess.run(
            [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', output_dir,
                word_path
            ],
            capture_output=True,
            timeout=60,
            text=True
        )
        
        # LibreOffice creates PDF with same name as input
        expected_pdf = os.path.join(
            output_dir,
            os.path.splitext(os.path.basename(word_path))[0] + '.pdf'
        )
        
        if os.path.exists(expected_pdf) and os.path.getsize(expected_pdf) > 0:
            # Rename to target path if different
            if expected_pdf != pdf_path:
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                os.rename(expected_pdf, pdf_path)
            print(f"SUCCESS: Converted using LibreOffice: {word_path} to {pdf_path}", file=sys.stderr)
            return True
        
    except Exception as e:
        print(f"Warning: LibreOffice conversion failed: {e}", file=sys.stderr)
    return False

# Method 3: Fallback to python-docx + reportlab (current implementation)
def try_reportlab(word_path: str, pdf_path: str) -> bool:
    """Fallback conversion using python-docx and reportlab"""
    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx library not installed. Install it with: pip install python-docx", file=sys.stderr)
        return False

    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    except ImportError:
        print("ERROR: reportlab library not installed. Install it with: pip install reportlab", file=sys.stderr)
        return False
    
    # Call the reportlab conversion function
    return convert_with_reportlab(word_path, pdf_path)

def convert_word_to_pdf(word_path: str, pdf_path: str) -> bool:
    """
    Convert Word document (DOCX) to PDF with perfect formatting preservation.
    Tries multiple methods in order of quality:
    1. docx2pdf (perfect match if Word/LibreOffice available)
    2. LibreOffice headless (excellent match, cross-platform)
    3. python-docx + reportlab (fallback with good formatting)
    
    Args:
        word_path: Path to input Word file (DOCX)
        pdf_path: Path to output PDF file
        
    Returns:
        True if conversion successful, False otherwise
    """
    # Check if Word file exists
    if not os.path.exists(word_path):
        print(f"ERROR: Word file not found: {word_path}", file=sys.stderr)
        return False
    
    # Try Method 1: docx2pdf (best quality - requires Word/LibreOffice)
    print("Attempting conversion with docx2pdf (best quality)...", file=sys.stderr)
    if try_docx2pdf(word_path, pdf_path):
        return True
    
    # Try Method 2: LibreOffice headless (excellent quality - cross-platform)
    print("Attempting conversion with LibreOffice headless (excellent quality)...", file=sys.stderr)
    if try_libreoffice(word_path, pdf_path):
        return True
    
    # Try Method 3: python-docx + reportlab (fallback)
    print("Using fallback method: python-docx + reportlab...", file=sys.stderr)
    return try_reportlab(word_path, pdf_path)

def convert_with_reportlab(word_path: str, pdf_path: str) -> bool:
    """
    Fallback conversion using python-docx and reportlab.
    This preserves good formatting but may not match exactly.
    """
    try:
        # Open Word document
        doc = Document(word_path)
        
        # Create PDF document
        pdf_doc = SimpleDocTemplate(
            pdf_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Build PDF content
        story = []
        styles = getSampleStyleSheet()
        
        # Use better fonts - Times-Roman for body text (more readable), Helvetica for headings
        # Try to register better fonts if available
        try:
            # Use Times-Roman as default (better for body text)
            default_font = 'Times-Roman'
            heading_font = 'Helvetica-Bold'
        except:
            default_font = 'Helvetica'
            heading_font = 'Helvetica-Bold'
        
        # Custom styles for headings with better fonts
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontName=heading_font,
            fontSize=24,
            textColor=colors.HexColor('#000000'),
            spaceAfter=12,
            spaceBefore=12,
            leading=28,  # Line spacing
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontName=heading_font,
            fontSize=20,
            textColor=colors.HexColor('#000000'),
            spaceAfter=10,
            spaceBefore=10,
            leading=24,
        )
        
        heading3_style = ParagraphStyle(
            'CustomHeading3',
            parent=styles['Heading3'],
            fontName=heading_font,
            fontSize=18,
            textColor=colors.HexColor('#000000'),
            spaceAfter=8,
            spaceBefore=8,
            leading=22,
        )
        
        # Helper function to convert RGB to hex color
        def rgb_to_hex(rgb):
            if rgb is None:
                return None
            if isinstance(rgb, str):
                return rgb
            # RGB is a tuple (R, G, B) with values 0-255
            return '#%02x%02x%02x' % (rgb[0], rgb[1], rgb[2])
        
        # Helper function to map Word font names to reportlab font names
        def map_font_name(word_font_name):
            """Map Word font names to reportlab-compatible font names"""
            if not word_font_name:
                return default_font
            
            word_font_lower = word_font_name.lower()
            
            # Map common font families
            font_mapping = {
                'times': 'Times-Roman',
                'times new roman': 'Times-Roman',
                'arial': 'Helvetica',
                'helvetica': 'Helvetica',
                'courier': 'Courier',
                'courier new': 'Courier',
                'calibri': 'Helvetica',  # Calibri not available, use Helvetica
                'verdana': 'Helvetica',   # Verdana not available, use Helvetica
                'tahoma': 'Helvetica',   # Tahoma not available, use Helvetica
            }
            
            for key, value in font_mapping.items():
                if key in word_font_lower:
                    return value
            
            # Default to Times-Roman for body text, Helvetica for headings
            return default_font
        
        # Helper function to build rich text with formatting
        def build_rich_text(paragraph):
            """Build rich text string with formatting tags for reportlab"""
            rich_text = []
            
            for run in paragraph.runs:
                text = run.text
                if not text:
                    continue
                
                # Escape XML special characters
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                
                # Build formatting tags
                tags = []
                
                # Bold
                if run.bold:
                    tags.append('b')
                
                # Italic
                if run.italic:
                    tags.append('i')
                
                # Font name
                font_name = None
                if run.font.name:
                    font_name = map_font_name(run.font.name)
                
                # Font color
                color = None
                if run.font.color and run.font.color.rgb:
                    color = rgb_to_hex(run.font.color.rgb)
                
                # Font size
                font_size = None
                if run.font.size:
                    # Convert from points (docx uses twentieths of a point)
                    font_size = run.font.size.pt if hasattr(run.font.size, 'pt') else None
                
                # Apply formatting
                formatted_text = text
                
                # Note: Font name is handled at paragraph style level, not in XML tags
                # Reportlab doesn't support inline font name changes easily
                # We'll use the default font for the paragraph style
                
                if tags:
                    for tag in reversed(tags):  # Close tags in reverse order
                        formatted_text = f'</{tag}>{formatted_text}'
                    for tag in tags:  # Open tags in order
                        formatted_text = f'<{tag}>{formatted_text}'
                
                # Apply color if present
                if color:
                    formatted_text = f'<font color="{color}">{formatted_text}</font>'
                
                # Apply font size if different from default
                if font_size and font_size != 12:
                    formatted_text = f'<font size="{int(font_size)}">{formatted_text}</font>'
                
                rich_text.append(formatted_text)
            
            return ''.join(rich_text) if rich_text else paragraph.text
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            
            if not text:
                story.append(Spacer(1, 6))
                continue
            
            # Determine style based on paragraph style
            style_name = paragraph.style.name.lower()
            
            # Build rich text with formatting
            rich_text = build_rich_text(paragraph)
            
            # Check paragraph alignment
            alignment = TA_LEFT
            if paragraph.alignment:
                if paragraph.alignment == 1:  # CENTER
                    alignment = TA_CENTER
                elif paragraph.alignment == 2:  # RIGHT
                    alignment = TA_RIGHT
                elif paragraph.alignment == 3:  # JUSTIFY
                    alignment = TA_JUSTIFY
            
            if 'heading 1' in style_name or 'title' in style_name:
                # Custom heading style with alignment and better font
                heading_style = ParagraphStyle(
                    'CustomHeading1',
                    parent=styles['Heading1'],
                    fontName=heading_font,
                    fontSize=24,
                    textColor=colors.HexColor('#000000'),
                    spaceAfter=12,
                    spaceBefore=12,
                    alignment=alignment,
                    leading=28,  # Line spacing (1.2x font size)
                )
                para = Paragraph(rich_text, heading_style)
            elif 'heading 2' in style_name:
                heading_style = ParagraphStyle(
                    'CustomHeading2',
                    parent=styles['Heading2'],
                    fontName=heading_font,
                    fontSize=20,
                    textColor=colors.HexColor('#000000'),
                    spaceAfter=10,
                    spaceBefore=10,
                    alignment=alignment,
                    leading=24,
                )
                para = Paragraph(rich_text, heading_style)
            elif 'heading 3' in style_name:
                heading_style = ParagraphStyle(
                    'CustomHeading3',
                    parent=styles['Heading3'],
                    fontName=heading_font,
                    fontSize=18,
                    textColor=colors.HexColor('#000000'),
                    spaceAfter=8,
                    spaceBefore=8,
                    alignment=alignment,
                    leading=22,
                )
                para = Paragraph(rich_text, heading_style)
            elif 'heading' in style_name:
                heading_style = ParagraphStyle(
                    'CustomHeading',
                    parent=styles['Heading3'],
                    fontName=heading_font,
                    fontSize=16,
                    textColor=colors.HexColor('#000000'),
                    spaceAfter=8,
                    spaceBefore=8,
                    alignment=alignment,
                    leading=20,
                )
                para = Paragraph(rich_text, heading_style)
            else:
                # Regular paragraph with formatting, alignment, and better font
                normal_style = ParagraphStyle(
                    'CustomNormal',
                    parent=styles['Normal'],
                    fontName=default_font,
                    fontSize=12,
                    textColor=colors.HexColor('#000000'),
                    spaceAfter=6,
                    alignment=alignment,
                    leading=14.4,  # 1.2x line spacing for better readability
                )
                para = Paragraph(rich_text, normal_style)
            
            story.append(para)
            story.append(Spacer(1, 6))
        
        # Process tables
        for table in doc.tables:
            # Extract table data with formatting
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Build rich text for each cell
                    cell_rich_text = []
                    for paragraph in cell.paragraphs:
                        para_text = build_rich_text(paragraph)
                        if para_text:
                            cell_rich_text.append(para_text)
                    cell_text = '<br/>'.join(cell_rich_text) if cell_rich_text else cell.text.strip()
                    row_data.append(Paragraph(cell_text, styles['Normal']) if cell_text else '')
                table_data.append(row_data)
            
            if table_data:
                # Create PDF table
                pdf_table = Table(table_data)
                
                # Determine if first row is header (usually bold or different style)
                is_header_row = False
                if len(table_data) > 0:
                    first_row = table.rows[0]
                    # Check if first row cells have bold formatting
                    for cell in first_row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.bold:
                                    is_header_row = True
                                    break
                            if is_header_row:
                                break
                        if is_header_row:
                            break
                
                table_style = [
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTSIZE', (0, 0), (-1, -1), 11),  # Slightly larger for better readability
                    ('FONTNAME', (0, 1), (-1, -1), default_font),  # Use Times-Roman for table body
                ]
                
                if is_header_row and len(table_data) > 0:
                    table_style.extend([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#366092')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('FONTNAME', (0, 0), (-1, 0), heading_font),  # Use Helvetica-Bold for headers
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('TOPPADDING', (0, 0), (-1, 0), 12),
                    ])
                
                table_style.extend([
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F5F5F5')]),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 1), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                ])
                
                pdf_table.setStyle(TableStyle(table_style))
                
                story.append(Spacer(1, 12))
                story.append(pdf_table)
                story.append(Spacer(1, 12))
        
        # Build PDF
        pdf_doc.build(story)
        
        # Check if output file was created
        if os.path.exists(pdf_path):
            print(f"SUCCESS: Converted {word_path} to {pdf_path}", file=sys.stderr)
            return True
        else:
            print(f"ERROR: Output file was not created: {pdf_path}", file=sys.stderr)
            return False
            
    except Exception as e:
        print(f"ERROR: Conversion failed: {str(e)}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python word_to_pdf.py <input_docx> <output_pdf>", file=sys.stderr)
        sys.exit(1)
    
    input_word = sys.argv[1]
    output_pdf = sys.argv[2]
    
    success = convert_word_to_pdf(input_word, output_pdf)
    sys.exit(0 if success else 1)

